#!/bin/python3

'''
###############################################################################
    __  ___          __  _                __  ____             __           
   /  |/  /__  ___  / /_(_)___  ____ _   /  |/  (_)___  __  __/ /____  _____
  / /|_/ / _ \/ _ \/ __/ / __ \/ __ `/  / /|_/ / / __ \/ / / / __/ _ \/ ___/
 / /  / /  __/  __/ /_/ / / / / /_/ /  / /  / / / / / / /_/ / /_/  __(__  ) 
/_/  /_/\___/\___/\__/_/_/ /_/\__, /  /_/  /_/_/_/ /_/\__,_/\__/\___/____/  
                             /____/                                         
###############################################################################


~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~
  DESCRIPTION  ######################
~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~

meeting_minutes.py utilizes OpenAI's Davinci model to take recorded speech and 
transcribe the recording into text. It then takes the text and uses OpenAI's
Davinci model to pull out the highlights of what was said during the meeting.
The script then takes the transcription and summary and writes them to a word
document. Finally, the document is timestamped.

~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~
  REQUIREMENTS  #####################
~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~
  [1]	python3
  [2]	Audacity (or other software to save audio recordings in .wav format)
  [3]	openai				-	python3 module
  [4]	docx				-	python3 module
  [5]	python-docx			-	python3 module
  [6]	speechrecognition	-	python3 module
'''

import openai
import docx
import speech_recognition as sr
import argparse
from datetime import datetime

# Define command line arguments
parser = argparse.ArgumentParser(description='Transcribe and summarize business meeting audio.')
parser.add_argument('audio_file', type=str, help='path to audio file in WAV format')
parser.add_argument('--api_key', type=str, help='OpenAI API key')

# Parse command line arguments
args = parser.parse_args()

# Set up OpenAI API key
if args.api_key:
    openai.api_key = args.api_key
else:
    openai.api_key = "YOUR_API_KEY"

# Set up speech recognition object
r = sr.Recognizer()

# Load audio file
with sr.AudioFile(args.audio_file) as source:
    audio_data = r.record(source)

# Transcribe audio using OpenAI Speech-to-Text API
transcription = r.recognize_google(audio_data)

# Split transcription into multiple parts
transcription_parts = [transcription[i:i+1000] for i in range(0, len(transcription), 1000)]

# Use ChatGPT to summarize meeting topics
summary_parts = []
for part in transcription_parts:
    summary = openai.Completion.create(
        engine="davinci",
        prompt="Summarize the topics discussed during the meeting:\n\n" + part,
        max_tokens=1024,
        n=1,
        stop=None,
        temperature=0.5,
    )
    summary_parts.append(summary.choices[0].text)

# Combine summary parts into final summary
summary = "\n".join(summary_parts)

# Save transcription and summary to Word document
timestamp = datetime.now().strftime("%Y%m%d-%H%M%S")
filename = f"meeting_minutes_{timestamp}.docx"
doc = docx.Document()
doc.add_paragraph("Meeting Timestamp: " + datetime.now().strftime("%m/%d/%Y %H:%M:%S"))
doc.add_paragraph("\nMeeting Transcription:\n\n" + transcription)
doc.add_paragraph("\nMeeting Summary:\n\n" + summary)
doc.save(filename)
print(f"Meeting minutes saved to {filename}.")

