#!/bin/python3

'''
###############################################################################
    __  ___          __  _                __  ____             __           
   /  |/  /__  ___  / /_(_)___  ____ _   /  |/  (_)___  __  __/ /____  _____
  / /|_/ / _ \/ _ \/ __/ / __ \/ __ `/  / /|_/ / / __ \/ / / / __/ _ \/ ___/
 / /  / /  __/  __/ /_/ / / / / /_/ /  / /  / / / / / / /_/ / /_/  __(__  ) 
/_/  /_/\___/\___/\__/_/_/ /_/\__, /  /_/  /_/_/_/ /_/\__,_/\__/\___/____/  
                             /____/                                         
###############################################################################


~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~
  DESCRIPTION  ######################
~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~

meeting_minutes.py utilizes OpenAI's Davinci model to take recorded speech and 
transcribe the recording into text. It then takes the text and uses OpenAI's
Davinci model to pull out the highlights of what was said during the meeting.
The script then takes the transcription and summary and writes them to a word
document. Finally, the document is timestamped.

~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~
  REQUIREMENTS  #####################
~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~
  [1]	python3
  [2]	Audacity (or other software to save audio recordings in .wav format)
  [3]	openai				-	python3 module
  [4]	docx				-	python3 module
  [5]	python-docx			-	python3 module
  [6]	speechrecognition	-	python3 module
'''

import argparse
from datetime import datetime
import os
import wave
import openai
import docx
import base64

def get_transcription(audio_file_path, api_key, chunk_size=30):
    openai.api_key = api_key
    
    # Get the length of the audio file in seconds
    audio_length = os.path.getsize(audio_file_path) / (2 * 16000)
    
    # Determine the number of chunks to break the audio into
    num_chunks = int(audio_length / chunk_size) + 1
    
    # Initialize the transcription string
    transcription = ""
    
    # Transcribe each chunk
    for i in range(num_chunks):
        # Calculate the start and end times for the chunk
        start_time = i * chunk_size
        end_time = min((i + 1) * chunk_size, audio_length)
        
        # Load the audio data for the chunk
        with open(audio_file_path, 'rb') as f:
            f.seek(int(start_time * 2 * 16000))
            audio_data = base64.b64encode(f.read(int((end_time - start_time) * 2 * 16000))).decode()
        
        # Send request to OpenAI API
        response = openai.Completion.create(
            engine='text-davinci-002',
            prompt=f'Transcribe the following audio file:\n{audio_data}\n',
            max_tokens=2048,
            temperature=0.5
        )
        
        # Add the transcribed text to the transcription string
        transcription += response.choices[0].text.strip()
        
        # Pause for a second to avoid hitting the rate limit
        time.sleep(1)
    
    print(f'[{datetime.now().strftime("%Y-%m-%d::%H%M.%S")}] Audio transcribed')
    return transcription

def generate_summary(transcription, api_key):
    openai.api_key = api_key

    # Split prompt into smaller chunks
    prompt_chunks = [transcription[i:i+2048] for i in range(0, len(transcription), 2048)]

    # Generate summary for each chunk
    print(f'[{datetime.now().strftime("%Y-%m-%d::%H%M.%S")}] Starting audio summarization')
    summaries = []
    for chunk in prompt_chunks:
        response = openai.Completion.create(
            engine='davinci',
            prompt=f'Summarize the following meeting transcription:\n{chunk}\n',
            max_tokens=1024,
            temperature=0.5
        )
        summaries.append(response.choices[0].text.strip())

    # Concatenate generated summaries
    summary = ' '.join(summaries)

    print(f'[{datetime.now().strftime("%Y-%m-%d::%H%M.%S")}] Audio summarized')
    return summary


def save_meeting_minutes(transcription, summary, file_path):
    # Create new Word document
    document = docx.Document()

    # Add meeting timestamp to document
    current_time = datetime.now().strftime('%m/%d/%Y %H:%M:%S')
    document.add_paragraph(f"Meeting Timestamp: {current_time}\n\n")

    # Add meeting transcription to document
    document.add_heading("Meeting Transcription:")
    document.add_paragraph(transcription)

    # Add meeting summary to document
    document.add_heading("Meeting Summary:")
    document.add_paragraph(summary)

    # Save document with timestamp in file name
    file_name = f"meeting_minutes_{current_time.replace('/', '_').replace(' ', '_').replace(':', '')}.docx"
    document.save(file_path + file_name)


if __name__ == '__main__':
    parser = argparse.ArgumentParser(description='Record business meetings and generate meeting minutes using OpenAI')
    parser.add_argument('--api_key', type=str, required=True, help='OpenAI API key')
    parser.add_argument('audio_file_path', type=str, help='Path to audio file')
    parser.add_argument('output_path', type=str, help='Output path for meeting minutes')
    args = parser.parse_args()

    # Get transcription from audio file
    transcription = get_transcription(args.audio_file_path, args.api_key, chunk_size=1)

    # Generate meeting summary from transcription
    summary = generate_summary(transcription, args.api_key)

    # Save meeting minutes to Word document
    print(f'[{datetime.now().strftime("%Y-%m-%d::%H%M.%S")}] Writing content to word document')
    save_meeting_minutes(transcription, summary, args.output_path)
    print(f'[{datetime.now().strftime("%Y-%m-%d::%H%M.%S")}] Meeting minutes saved to {file_name}')

